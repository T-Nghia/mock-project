import tempfile
import unittest
from pathlib import Path
from unittest.mock import MagicMock, patch

from docx import Document as DocxDocument

from app.utils.text_extract import extract_text


class TextExtractTestCase(unittest.TestCase):
    def save_docx(self, build_document) -> tuple[tempfile.TemporaryDirectory, Path]:
        temp_dir = tempfile.TemporaryDirectory()
        path = Path(temp_dir.name) / "lesson.docx"
        document = DocxDocument()
        build_document(document)
        document.save(path)
        return temp_dir, path

    def test_txt_preserves_line_breaks_and_trims_outer_whitespace(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "lesson.txt"
            path.write_text("  Dong mot\nDong hai  \n", encoding="utf-8")

            result = extract_text(str(path), "txt")

        self.assertEqual(result, "Dong mot\nDong hai")

    @patch("app.utils.text_extract.PdfReader")
    def test_pdf_joins_non_empty_pages_with_blank_lines(self, pdf_reader):
        first_page = MagicMock()
        first_page.extract_text.return_value = "  Trang mot  "
        empty_page = MagicMock()
        empty_page.extract_text.return_value = "   "
        last_page = MagicMock()
        last_page.extract_text.return_value = "Trang ba"
        pdf_reader.return_value.pages = [first_page, empty_page, last_page]

        result = extract_text("lesson.pdf", "pdf")

        self.assertEqual(result, "Trang mot\n\nTrang ba")
        pdf_reader.assert_called_once_with("lesson.pdf")

    def test_unsupported_format_logs_warning_and_returns_empty_string(self):
        with self.assertLogs("app.utils.text_extract", level="WARNING") as logs:
            result = extract_text("slides.pptx", "pptx")

        self.assertEqual(result, "")
        self.assertIn("Unsupported extraction format", logs.output[0])
        self.assertIn("pptx", logs.output[0])
        self.assertIn("slides.pptx", logs.output[0])

    @patch("app.utils.text_extract.PdfReader", side_effect=ValueError("invalid pdf"))
    def test_supported_format_error_logs_exception_and_returns_empty_string(self, _pdf_reader):
        with self.assertLogs("app.utils.text_extract", level="ERROR") as logs:
            result = extract_text("broken.pdf", "pdf")

        self.assertEqual(result, "")
        self.assertIn("Text extraction failed", logs.output[0])
        self.assertIn("pdf", logs.output[0])
        self.assertIn("broken.pdf", logs.output[0])

    def test_docx_extracts_non_empty_paragraphs(self):
        def build(document):
            document.add_paragraph("  Gioi thieu  ")
            document.add_paragraph("")
            document.add_paragraph("Ket luan")

        temp_dir, path = self.save_docx(build)
        self.addCleanup(temp_dir.cleanup)

        result = extract_text(str(path), "docx")

        self.assertEqual(result, "Gioi thieu\n\nKet luan")

    def test_docx_formats_table_rows_as_header_value_pairs(self):
        def build(document):
            table = document.add_table(rows=3, cols=3)
            headers = ["Mon hoc", "So tin chi", "Diem"]
            first_row = ["Python", "3", "8.5"]
            second_row = ["Database", "4", ""]
            for index, value in enumerate(headers):
                table.rows[0].cells[index].text = value
            for index, value in enumerate(first_row):
                table.rows[1].cells[index].text = value
            for index, value in enumerate(second_row):
                table.rows[2].cells[index].text = value

        temp_dir, path = self.save_docx(build)
        self.addCleanup(temp_dir.cleanup)

        result = extract_text(str(path), "docx")

        self.assertEqual(
            result,
            "[TABLE 1]\n"
            "Mon hoc: Python; So tin chi: 3; Diem: 8.5\n"
            "Mon hoc: Database; So tin chi: 4",
        )

    def test_docx_preserves_paragraph_table_paragraph_order(self):
        def build(document):
            document.add_paragraph("Truoc bang")
            table = document.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = "Muc"
            table.rows[0].cells[1].text = "Gia tri"
            table.rows[1].cells[0].text = "A"
            table.rows[1].cells[1].text = "10"
            document.add_paragraph("Sau bang")

        temp_dir, path = self.save_docx(build)
        self.addCleanup(temp_dir.cleanup)

        result = extract_text(str(path), "docx")

        self.assertEqual(
            result,
            "Truoc bang\n\n[TABLE 1]\nMuc: A; Gia tri: 10\n\nSau bang",
        )

    def test_docx_empty_header_uses_column_fallback_and_header_only_table_is_ignored(self):
        def build(document):
            table = document.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = ""
            table.rows[0].cells[1].text = "Gia tri"
            table.rows[1].cells[0].text = "A"
            table.rows[1].cells[1].text = "10"
            header_only = document.add_table(rows=1, cols=1)
            header_only.rows[0].cells[0].text = "Khong co du lieu"

        temp_dir, path = self.save_docx(build)
        self.addCleanup(temp_dir.cleanup)

        result = extract_text(str(path), "docx")

        self.assertEqual(result, "[TABLE 1]\nCột 1: A; Gia tri: 10")


if __name__ == "__main__":
    unittest.main()
